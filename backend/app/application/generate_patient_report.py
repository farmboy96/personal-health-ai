from __future__ import annotations

import json
import logging
import math
import random
import re
from datetime import date, datetime
from pathlib import Path
from typing import Any

from sqlalchemy import desc, func

from app.ai.client import get_openai_client
from app.application._docx_helpers import (
    add_text_with_bullets,
    apply_table_layout,
    configure_document_typography,
    ensure_bullet_styles,
    set_cell_text,
    shade_cell,
    truncate,
)
from app.db.database import SessionLocal
from app.db.models import ClinicalReport, GeneticRecommendation, GeneticVariant, LabResult
from app.domain.assessment.biological_age import (
    compute_fitness_age,
    compute_framingham_heart_age,
    compute_phenotypic_age,
)

logger = logging.getLogger(__name__)

HEDGING_TERMS = ["consider", "may want to", "it might be beneficial"]
TIMELINE_RULES_BLOCK = """TIMELINE REASONING RULES (mandatory):
1. Labs are dated. Supplements and behaviors are dated. Never attribute a lab value to an intervention that started after the lab was drawn.
2. The most recent lab draw is January 27, 2026. If a supplement or behavior started AFTER January 27, 2026, it cannot have affected any current lab value. Do not write phrases like "consistent with [new intervention]" or "reflecting [new intervention]" for any current biomarker.
3. When writing about future lab trajectories, use explicit forward-looking framing: "Expected in July 2026 labs after TMG exposure" — not "consistent with your TMG".
4. When writing "What's working" sections, attribute improvements only to interventions that were in effect before the lab draw date. For recent metabolic wins (A1C, insulin, hs-CRP, weight), attribute to longstanding behaviors (fasting, sugar elimination since January 2026, BJJ, weight loss) — these were in effect before the draw.
5. For physiology rollups (sleep, HRV, steps) drawn from Apple Health, attribute freely to current behaviors because rollups are current by definition.
6. Blood donation: Robert has donated at OBI multiple times historically. Most recent donation before Jan 27 2026 draw was July 2025. New donation in April 2026. Hematocrit 52.6% on Jan 27 reflects ~6 months without donation. Current hematocrit is likely lower but unmeasured until next draw.
7. SEASONAL BEHAVIORS: Weighted-vest rucking only occurs April through October (tied to mowing acreage). For the January 27 2026 draw, rucking was NOT in effect. Do not attribute January labs to rucking. BJJ is year-round and is always in effect."""


def _default_output_dir() -> Path:
    d = Path(__file__).resolve().parents[2] / "data" / "reports"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _extract_pmids(text: str) -> list[str]:
    return re.findall(r"(?:\[?\s*PMID:\s*)(\d+)\]?", text or "")


def _dehedge_text(text: str) -> str:
    t = text or ""
    t = re.sub(r"\b[Cc]onsider\b", "Do", t)
    t = re.sub(r"\b[Mm]ay want to\b", "should", t)
    t = re.sub(r"\bit might be beneficial\b", "it is useful", t, flags=re.I)
    return t


def _normalize_genotype(gt: str) -> tuple[str, str]:
    cleaned = (gt or "").replace("(", "").replace(")", "").replace(" ", "")
    parts = cleaned.split(";")
    if len(parts) != 2:
        return "", ""
    return parts[0].upper(), parts[1].upper()


def _match_pattern(genotype: str, pattern: str) -> bool:
    a1, a2 = _normalize_genotype(genotype)
    if not a1:
        return False
    p = (pattern or "any").strip().lower()
    if p == "any":
        return True
    if p == "homozygous_minor":
        return a1 == a2
    if p.startswith("*/"):
        target = p.split("/", 1)[1].upper()
        return a1 == target or a2 == target
    if "/" in p:
        x, y = p.split("/", 1)
        x = x.upper()
        y = y.upper()
        return sorted([a1, a2]) == sorted([x, y])
    return False


def _load_report_row(db, clinical_report_id: int | None) -> ClinicalReport:
    if clinical_report_id is not None:
        row = db.query(ClinicalReport).filter(ClinicalReport.id == clinical_report_id).first()
    else:
        row = db.query(ClinicalReport).order_by(desc(ClinicalReport.created_at)).first()
    if not row:
        raise ValueError("No ClinicalReport found")
    return row


def _match_recommendations(db, variants: list[GeneticVariant]) -> list[dict[str, Any]]:
    var_by_rsid: dict[str, list[GeneticVariant]] = {}
    for v in variants:
        var_by_rsid.setdefault(v.rsid, []).append(v)

    recs = db.query(GeneticRecommendation).all()
    matched: list[dict[str, Any]] = []
    for r in recs:
        vrows = var_by_rsid.get(r.rsid, [])
        for v in vrows:
            if _match_pattern(v.genotype or "", r.genotype_pattern or "any"):
                d = {
                    "rsid": r.rsid,
                    "gene": r.gene,
                    "category": r.category,
                    "recommendation_text": r.recommendation_text,
                    "rationale": r.rationale,
                    "action_level": r.action_level,
                    "priority": r.priority,
                    "source_notes": r.source_notes,
                    "genotype": v.genotype,
                }
                matched.append(d)
                break

    # de-dupe exact text
    uniq = {}
    for m in matched:
        k = (m["rsid"], m["category"], m["recommendation_text"])
        if k not in uniq:
            uniq[k] = m
    out = list(uniq.values())
    out.sort(key=lambda x: (x["priority"], x["category"], x["recommendation_text"]))
    return out


def _extract_normalized_numbers(text: str) -> set[str]:
    raw = re.findall(r"\d{1,3}(?:,\d{3})*(?:\.\d+)?", text or "")
    out: set[str] = set()
    for tok in raw:
        out.add(tok.replace(",", ""))
    return out


def _numbers_from_context(snapshot: dict[str, Any]) -> set[str]:
    blob = json.dumps(snapshot, default=str)
    return _extract_normalized_numbers(blob)


def _timeline_context_from_snapshot(snapshot: dict[str, Any]) -> str:
    return (
        "Most recent lab draw date: 2026-01-27\n"
        "Known timeline notes:\n"
        "- Nordic Naturals Ultimate Omega 2X started 2026-04-20 (cannot affect Jan 2026 labs).\n"
        "- TMG 2000 mg/day started 2026-04-20 (cannot affect Jan 2026 labs).\n"
        "- Super K started 2026-04-20 (cannot affect Jan 2026 labs).\n"
        "- Previous fish oil formula was discontinued 2026-04-20.\n"
        "- Donation timeline: prior donation before Jan draw was July 2025; next donation April 2026.\n"
        "Current numeric context JSON:\n"
        + json.dumps(snapshot, indent=2, default=str)
    )


def validate_timeline_attribution(draft_text: str, timeline_context: str) -> dict[str, Any]:
    client = get_openai_client()
    prompt = (
        f"Timeline:\n{timeline_context}\n\n"
        f"Draft:\n{draft_text}\n\n"
        "Return strict JSON only with schema:\n"
        '{"violations":[{"line":"string","issue":"string","suggested_fix":"string"}],"clean":boolean}'
    )
    rsp = client.chat.completions.create(
        model="gpt-5",
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a fact-checker for clinical reports. Given the timeline and a draft section, "
                    "identify any causal or associative attribution that violates the timeline. Flag specifically "
                    "cases where a biomarker value is attributed to an intervention that started after the biomarker was measured. "
                    "Specifically flag any attribution of January 2026 labs to rucking, mowing, weighted vest, or weighted mowing."
                ),
            },
            {"role": "user", "content": prompt},
        ],
    )
    raw = (rsp.choices[0].message.content or "").strip()
    seasonal_re = re.compile(r"\b(rucking|mowing|weighted vest|weighted-vest|weighted mowing)\b", re.I)
    january_re = re.compile(r"\b(jan(?:uary)?\s*2026|january 27[, ]*2026|january labs|pre-?draw|draw)\b", re.I)
    try:
        data = json.loads(raw)
        if isinstance(data, dict):
            violations = data.get("violations") or []
            clean = bool(data.get("clean", False))
            # Deterministic guard for known seasonal attribution failure mode.
            if seasonal_re.search(draft_text) and january_re.search(draft_text):
                violations.append(
                    {
                        "line": "Detected January-lab attribution language near rucking/mowing/weighted-vest behavior.",
                        "issue": "Rucking is seasonal (April-October) and was not in effect for Jan 27, 2026 labs.",
                        "suggested_fix": "Remove causal attribution to rucking for January labs; use BJJ/year-round behaviors or forward-looking framing for July labs.",
                    }
                )
                clean = False
            return {
                "violations": violations,
                "clean": clean,
                "raw": raw,
            }
    except Exception:
        pass
    return {
        "violations": [],
        "clean": False,
        "raw": raw,
    }


def _rewrite_timeline_violations(
    *,
    draft_text: str,
    timeline_context: str,
    violations: list[dict[str, Any]],
) -> str:
    client = get_openai_client()
    prompt = (
        "Rewrite this draft to remove timeline-attribution violations while preserving all valid numeric facts.\n\n"
        f"Timeline:\n{timeline_context}\n\n"
        f"Violations:\n{json.dumps(violations, ensure_ascii=False, indent=2)}\n\n"
        f"Draft:\n{draft_text}\n\n"
        "Rules:\n"
        "- Keep tone and structure.\n"
        "- Do not invent numbers.\n"
        "- Use forward-looking wording for post-2026-01-27 interventions.\n"
    )
    rsp = client.chat.completions.create(
        model="gpt-5",
        messages=[
            {
                "role": "system",
                "content": "You edit patient report text to remove timeline-causality errors without adding new claims.",
            },
            {"role": "user", "content": prompt},
        ],
    )
    return (rsp.choices[0].message.content or "").strip()


def _ai_section1(patient_context_snapshot: dict[str, Any]) -> str:
    ctx = json.dumps(patient_context_snapshot, indent=2, default=str)
    prompt = f"""Write Section 1: \"How you're doing right now\" for Robert Grogan.

Use only this data:
{ctx}

Output 3-5 short paragraphs. Requirements:
- Name specific numbers and trajectories that are positive or improving.
- Tie each number to an observed behavior where plausible.
- Do not invent numbers.
- Tone: direct, informed, zero fluff.

{TIMELINE_RULES_BLOCK}
"""
    client = get_openai_client()
    rsp = client.chat.completions.create(
        model="gpt-5",
        messages=[
            {
                "role": "system",
                "content": "You write precise patient-facing health analysis from provided numbers only, with strict timeline causality discipline.",
            },
            {"role": "user", "content": prompt},
        ],
    )
    return (rsp.choices[0].message.content or "").strip()


def _ai_top_5_actions(
    *,
    patient_context_snapshot: dict[str, Any],
    matched_recommendations: list[dict[str, Any]],
    topic_narratives: dict[str, str],
) -> str:
    context_json = json.dumps(patient_context_snapshot, indent=2, default=str)
    recs_json = json.dumps(
        [
            {
                "category": r.get("category"),
                "action_level": r.get("action_level"),
                "priority": r.get("priority"),
                "recommendation_text": r.get("recommendation_text"),
                "rationale": r.get("rationale"),
                "rsid": r.get("rsid"),
                "gene": r.get("gene"),
            }
            for r in matched_recommendations
        ],
        ensure_ascii=False,
    )
    topics_json = json.dumps(topic_narratives, ensure_ascii=False)

    prompt = """
From the patient context, matched genetic recommendations, and topic narratives below, select the 5 highest-leverage actions for the patient to focus on over the next 30 days. Criteria:

1. Rank by expected impact on the patient's most concerning biomarker trends (rising LDL, elevated homocysteine, elevated estradiol)
2. Favor actions the patient can execute alone (diet swaps, routine changes) over discuss-with-doctor items
3. Each action must be phrased in plain English at an 8th grade reading level
4. Each action is 1 sentence describing the change, followed by 1-2 sentences explaining why it matters in terms the patient can feel (not medical jargon)
5. Order by priority, highest leverage first
6. No PMID citations in this section — keep it clean; citations live in the detail sections

Write in direct voice: "Switch X to Y" not "consider switching X to Y." The reader is motivated and capable; give him the action plainly.

Output exactly 5 top-level bullets, each beginning with "- ".

{TIMELINE_RULES_BLOCK}
"""
    full_prompt = (
        f"{prompt}\n\n"
        f"Patient context snapshot JSON:\n{context_json}\n\n"
        f"Matched recommendations JSON:\n{recs_json}\n\n"
        f"Topic narratives JSON:\n{topics_json}\n"
    )
    client = get_openai_client()
    rsp = client.chat.completions.create(
        model="gpt-5",
        messages=[
            {
                "role": "system",
                "content": "You create direct patient action plans. No hedging. No citations in this section. Follow strict timeline causality rules.",
            },
            {"role": "user", "content": full_prompt},
        ],
    )
    return _dehedge_text((rsp.choices[0].message.content or "").strip())


def _prioritize_august_handoff(questions_grouped: dict[str, list[str]], max_items: int = 6) -> list[str]:
    # Balanced handoff distribution (5-6 items):
    # - 1-2 lipid, 1 TRT, 1 homocysteine, 1 prostate, 1 monitoring/hematocrit
    topic_ldl = questions_grouped.get("Topic: ldl_trajectory_trt_apoa2", [])
    topic_mthfr = questions_grouped.get("Topic: mthfr_homocysteine_trt", [])
    topic_prostate = questions_grouped.get("Topic: prostate_surveillance_trt_high_prs", [])
    monitoring = questions_grouped.get("Monitoring", [])
    all_items: list[str] = []
    for _, items in questions_grouped.items():
        all_items.extend(items)

    def score(item: str) -> tuple[int, int]:
        t = item.lower()
        s = 0
        if any(k in t for k in ("ldl", "apob", "lipid", "cholesterol", "triglycer", "mpo", "lp-pla2")):
            s += 7
        if any(k in t for k in ("testosterone", "trt", "estradiol", "dose")):
            s += 6
        if any(k in t for k in ("homocysteine", "methyl", "folate", "b12")):
            s += 5
        if any(k in t for k in ("psa", "mri", "prostate", "biopsy", "urology")):
            s += 5
        if any(k in t for k in ("hematocrit", "ferritin", "donation", "phlebotomy", "cbc")):
            s += 4
        return (-s, len(item))

    selected: list[str] = []

    def _pick(items: list[str], keywords: tuple[str, ...]) -> str | None:
        ranked = sorted(set(items), key=score)
        for cand in ranked:
            if cand in selected:
                continue
            lc = cand.lower()
            if any(k in lc for k in keywords):
                return cand
        return None

    # Lipid 1-2
    lipid_keywords = ("ldl", "apob", "lipid", "cholesterol", "triglycer", "mpo", "lp-pla2")
    x = _pick(topic_ldl or all_items, lipid_keywords)
    if x:
        selected.append(x)
    x = _pick(topic_ldl or all_items, lipid_keywords)
    if x:
        selected.append(x)

    # TRT
    trt_keywords = ("testosterone", "trt", "estradiol", "dose")
    x = _pick(topic_mthfr + topic_ldl + all_items, trt_keywords)
    if x:
        selected.append(x)

    # Homocysteine
    homocys_keywords = ("homocysteine", "methyl", "folate", "b12")
    x = _pick(topic_mthfr + all_items, homocys_keywords)
    if x:
        selected.append(x)

    # Prostate
    prostate_keywords = ("psa", "mri", "prostate", "biopsy", "urology")
    x = _pick(topic_prostate + all_items, prostate_keywords)
    if x:
        selected.append(x)

    # Monitoring / hematocrit
    monitoring_keywords = ("hematocrit", "ferritin", "donation", "phlebotomy", "cbc")
    x = _pick(monitoring + all_items, monitoring_keywords)
    if x:
        selected.append(x)

    # Fill remaining slots by urgency.
    for cand in sorted(set(all_items), key=score):
        if cand in selected:
            continue
        selected.append(cand)
        if len(selected) >= max_items:
            break

    return selected[:max_items]


def _topic_translation_ai(
    *,
    topic_title: str,
    physician_narrative: str,
    patient_context: str,
    allowed_studies: list[dict[str, Any]],
    obi_ldl_addon: str = "",
    topic_key: str | None = None,
) -> str:
    allowed_pmids = [str(s.get("pmid")) for s in allowed_studies if s.get("pmid")]
    studies_blob = json.dumps(
        [
            {
                "pmid": s.get("pmid"),
                "title": s.get("title"),
                "journal": s.get("journal"),
                "year": s.get("year"),
                "abstract": s.get("abstract"),
                "summary": s.get("summary"),
                "applicability_note": s.get("applicability_note"),
            }
            for s in allowed_studies
        ],
        ensure_ascii=False,
    )
    prompt = f"""
You are translating a physician-facing clinical literature synthesis into a patient-facing action guide for a specific reader. The reader is a 54-year-old man, a 3-stripe brown belt in Brazilian jiu-jitsu, who does weighted-vest rucking on weekends, practices 16:8 intermittent fasting plus occasional multi-day fasts, and works from home at a computer. He reads Peter Attia, Andrew Huberman, and Rhonda Patrick. He can handle real information.

Your voice: direct, evidence-citing, action-oriented. No hedging for its own sake. Do not use phrases like "consider," "may want to," or "it might be beneficial." Prefer imperatives: "swap X for Y," "add X," "stop doing Y." When evidence is strong, say so plainly. When it's weak, say so plainly — do not pretend certainty you don't have.

You may only cite PMIDs provided in the context. Do not invent studies, authors, effect sizes, or dates. If you reference a number from a study, it must appear verbatim in the provided abstracts.

Structure each topic as four sections with these exact headings:

What's working — specific things in the patient's data or behavior that are moving numbers in the right direction. Be specific — name the number and name the behavior.

What to adjust — diet, sleep, training, behavior changes the patient can make alone. Name specific swaps. Cite PMIDs when the adjustment is evidence-backed.

What to discuss with Dr. Lamkin — supplement additions or dose changes, new labs to order, medication considerations. These are the handoff items.

What to watch for — specific numbers or symptoms that should trigger contacting Dr. Lamkin earlier than the scheduled visit.

Length: each section 2-5 bullet points. Do not pad.

Topic: {topic_title}
Patient context:
{patient_context}

Additional longitudinal context (patient-provided screening / OBI app — use when relevant):
{obi_ldl_addon if obi_ldl_addon.strip() else "(none)"}

Physician narrative:
{physician_narrative}

Allowed studies JSON:
{studies_blob}

Allowed PMIDs only: {", ".join(allowed_pmids)}

{TIMELINE_RULES_BLOCK}
"""
    if "ldl_trajectory_trt_apoa2" in topic_title.lower() or "ldl trajectory" in topic_title.lower():
        prompt += """
Additional required instruction for this LDL topic:
- The patient's January 27 2026 labs include two specialty cardiovascular inflammation markers — MPO 573 pmol/L (optimal <470) and Lp-PLA2 activity 209 nmol/min/mL (optimal <124), both flagged high.
- Reference these in the "What's working" or "What to watch for" sections as supporting evidence for the urgency of lipid management.
- Do not invent effect sizes; just report the values and cite the flagging.
"""
        if obi_ldl_addon.strip():
            prompt += """
OBI screening cholesterol trajectory (finger-stick at donation; distinct from venous lab TOTAL_CHOLESTEROL):
- When the additional longitudinal context block lists dated OBI total cholesterol values, incorporate the multi-year trajectory into "What's working" and/or "What to adjust" where it strengthens the narrative (directionally reliable; less precise than venous labs).
- The March 28, 2026 OBI total cholesterol (311 mg/dL in the patient-provided series) is the highest on record in that dataset; cholesterol has been chronically elevated since at least September 2022 in that series.
- In "What to discuss with Dr. Lamkin", explicitly include pulling the statin / intensive LDL-lowering pharmacotherapy conversation forward — grounded in both venous LDL from January 2026 labs and this multi-year OBI screening trajectory with the March 2026 peak.
"""
    if topic_key == "prostate_surveillance_trt_high_prs":
        prompt += """
Additional required instruction for PSA wording (prostate surveillance topic):
- When describing the patient's PSA, use the actual numeric value from the labs. Never describe a measured PSA value as "undetectable" — that term means below assay detection limit (typically <0.1 ng/mL). For PSA values within reference range but >0.1, use phrases like "within normal range at X ng/mL" or "low-normal at X ng/mL" depending on the value.
"""
    client = get_openai_client()
    rsp = client.chat.completions.create(
        model="gpt-5",
        messages=[
            {
                "role": "system",
                "content": "You produce direct patient action guides with strict PMID discipline, no invented claims, and strict timeline causality discipline.",
            },
            {"role": "user", "content": prompt},
        ],
    )
    text = (rsp.choices[0].message.content or "").strip()
    # hard guard against hallucinated PMIDs
    bad = [p for p in _extract_pmids(text) if p not in set(allowed_pmids)]
    if bad:
        for p in bad:
            text = text.replace(f"[PMID: {p}]", "[PMID: removed]")
            text = text.replace(f"PMID: {p}", "PMID: removed")
    return _dehedge_text(text)


def _extract_discuss_items_from_topic(text: str) -> list[str]:
    out: list[str] = []
    lines = (text or "").splitlines()
    in_section = False
    for ln in lines:
        t = ln.strip()
        if not t:
            continue
        if t.lower().startswith("what to discuss with dr. lamkin"):
            in_section = True
            continue
        if in_section and t.lower().startswith("what to watch for"):
            break
        if in_section and t.startswith("- "):
            out.append(t[2:].strip())
    return out


def _category_display(cat: str) -> str:
    m = {
        "diet": "Diet",
        "exercise": "Exercise",
        "supplement": "Supplements",
        "sleep": "Sleep & Lifestyle",
        "behavior": "Sleep & Lifestyle",
        "monitoring": "Monitoring",
    }
    return m.get(cat, cat.title())


def _fmt_bio_metric_value(d: dict[str, Any]) -> str:
    if not d.get("computable"):
        return "—"
    v = float(d.get("value") or float("nan"))
    if math.isnan(v):
        return "—"
    return f"{v:.1f}".rstrip("0").rstrip(".")


def _fmt_bio_delta_years(d: dict[str, Any]) -> str:
    if not d.get("computable"):
        return "—"
    x = float(d.get("delta") or float("nan"))
    if math.isnan(x):
        return "—"
    sign = "+" if x > 0 else ""
    return f"{sign}{x:.1f} years"


def _interpret_phenotypic_row(d: dict[str, Any]) -> str:
    if not d.get("computable"):
        return "More labs needed for Levine phenotypic age."
    delta = float(d.get("delta") or 0.0)
    if delta <= -3:
        return "You appear metabolically younger than your calendar age."
    if delta >= 3:
        return "This estimate suggests faster biological aging than calendar age."
    return "Roughly in line with your calendar age."


def _interpret_fitness_row(d: dict[str, Any]) -> str:
    if not d.get("computable"):
        return "Resting HR / BMI inputs were insufficient for this estimate."
    delta = float(d.get("delta") or 0.0)
    if delta <= -3:
        return "Estimated cardio fitness is stronger than typical for your age."
    if delta >= 3:
        return "Estimated fitness is below typical for your age."
    return "Estimated fitness is close to typical for your age."


def _interpret_framingham_row(d: dict[str, Any]) -> str:
    if not d.get("computable"):
        return "Total cholesterol / HDL needed for heart age."
    delta = float(d.get("delta") or 0.0)
    if delta <= -3:
        return "Estimated vascular risk profile is younger than calendar age."
    return "Your lipid profile contributes to elevated estimated cardiovascular risk vs calendar age."


def _add_biological_age_snapshot_docx(doc, phen: dict[str, Any], fit: dict[str, Any], fram: dict[str, Any]) -> None:
    from docx.shared import Pt

    doc.add_heading("Your biological age snapshot", level=2)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    intro.add_run(
        "Biological age metrics are different views of how your body is aging compared to calendar time. "
        "Each number answers a different question; none is a single 'truth.'"
    )

    chrono = phen.get("chronological_age") or fit.get("chronological_age") or fram.get("chronological_age")
    try:
        chrono_s = str(int(round(float(chrono)))) if chrono is not None else "—"
    except (TypeError, ValueError):
        chrono_s = "—"

    tbl = doc.add_table(rows=4, cols=5)
    apply_table_layout(tbl, [1.45, 0.95, 1.05, 1.15, 2.45])
    headers = ["Metric", "Your age", "Chronological", "Delta", "Interpretation"]
    for i, h in enumerate(headers):
        set_cell_text(tbl.rows[0].cells[i], h, bold=True, size_pt=9)
        shade_cell(tbl.rows[0].cells[i], "D9D9D9")

    rows_spec = [
        ("Phenotypic Age", phen, _interpret_phenotypic_row),
        ("Fitness Age", fit, _interpret_fitness_row),
        ("Framingham Heart Age", fram, _interpret_framingham_row),
    ]
    for ri, (label, pack, interp_fn) in enumerate(rows_spec, start=1):
        cells = tbl.rows[ri].cells
        set_cell_text(cells[0], label, bold=False, size_pt=9)
        set_cell_text(cells[1], _fmt_bio_metric_value(pack), size_pt=9)
        set_cell_text(cells[2], chrono_s, size_pt=9)
        set_cell_text(cells[3], _fmt_bio_delta_years(pack), size_pt=9)
        set_cell_text(cells[4], interp_fn(pack), size_pt=9)

    if phen.get("computable") and not math.isnan(float(phen.get("delta") or float("nan"))) and float(phen["delta"]) > 3:
        ca = doc.add_paragraph()
        ca.paragraph_format.space_before = Pt(6)
        ca.paragraph_format.space_after = Pt(6)
        ca.add_run(
            "Phenotypic age is notably higher than calendar age on this snapshot — driven by the biomarker "
            "pattern (inflammation, organ markers, blood counts) rather than lipids alone."
        )

    if fit.get("computable") and not math.isnan(float(fit.get("delta") or float("nan"))) and float(fit["delta"]) > 3:
        ca = doc.add_paragraph()
        ca.paragraph_format.space_before = Pt(6)
        ca.paragraph_format.space_after = Pt(6)
        ca.add_run(
            "Fitness age reads older than calendar age — worth interpreting alongside resting heart rate trends "
            "and training load; this is an approximation, not a treadmill test."
        )

    if fram.get("computable") and not math.isnan(float(fram.get("delta") or float("nan"))) and float(fram["delta"]) > 3:
        ca = doc.add_paragraph()
        ca.paragraph_format.space_before = Pt(6)
        ca.paragraph_format.space_after = Pt(6)
        ca.add_run(
            "Framingham heart age is elevated versus calendar age — primarily reflecting your lipid profile "
            "(high total cholesterol and comparatively low HDL), which increases estimated 10-year general "
            "cardiovascular risk."
        )


def generate_patient_report(clinical_report_id: int | None = None, output_dir: str | None = None) -> dict:
    db = SessionLocal()
    try:
        report_row = _load_report_row(db, clinical_report_id)
        snapshot = json.loads(report_row.patient_context_snapshot or "{}")
        narratives = json.loads(report_row.narrative_sections or "{}")
        retrieved = json.loads(report_row.retrieved_studies or "[]")
        variants = db.query(GeneticVariant).all()
        matched = _match_recommendations(db, variants)
    finally:
        db.close()

    out_dir = Path(output_dir) if output_dir else _default_output_dir()
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    out_path = out_dir / f"patient_report_{ts}.docx"

    patient_context = json.dumps(snapshot, default=str, indent=2)
    timeline_context = _timeline_context_from_snapshot(snapshot)
    section1_initial = _dehedge_text(_ai_section1(snapshot))
    section1_validation = validate_timeline_attribution(section1_initial, timeline_context)
    section1 = section1_initial
    if not section1_validation.get("clean") and section1_validation.get("violations"):
        section1 = _dehedge_text(
            _rewrite_timeline_violations(
                draft_text=section1_initial,
                timeline_context=timeline_context,
                violations=section1_validation["violations"],
            )
        )

    # translate each topic
    topic_translations: dict[str, str] = {}
    topic_initials: dict[str, str] = {}
    topic_validation: dict[str, dict[str, Any]] = {}
    cited_pmids: set[str] = set()
    by_topic_studies: dict[str, list[dict[str, Any]]] = {}
    for s in retrieved:
        by_topic_studies.setdefault(str(s.get("topic_key") or ""), []).append(s)

    from app.domain.assessment.obi_screening import (
        format_obi_bp_line_for_prompt,
        format_obi_cholesterol_trajectory_for_prompt,
    )

    db_obi = SessionLocal()
    try:
        mx_lab = db_obi.query(func.max(LabResult.lab_date)).scalar()
        ref_lab_date: date = mx_lab if isinstance(mx_lab, date) else date(2026, 1, 27)
        obi_chol_blob = format_obi_cholesterol_trajectory_for_prompt(db_obi)
        obi_bp_line = format_obi_bp_line_for_prompt(db_obi, ref_lab_date)
    finally:
        db_obi.close()

    obi_ldl_parts: list[str] = []
    if obi_chol_blob.strip():
        obi_ldl_parts.append(obi_chol_blob.strip())
    if obi_bp_line.strip():
        obi_ldl_parts.append(obi_bp_line.strip())
    obi_ldl_combined = "\n\n".join(obi_ldl_parts)

    for topic_key, physician_text in narratives.items():
        topic_title = topic_key
        if topic_key in {
            "mthfr_homocysteine_trt",
            "ldl_trajectory_trt_apoa2",
            "prostate_surveillance_trt_high_prs",
        }:
            from app.research.topic_catalog import TOPICS

            topic_title = TOPICS[topic_key]["title"]
        allowed = by_topic_studies.get(topic_key, [])
        translated = _topic_translation_ai(
            topic_title=topic_title,
            physician_narrative=physician_text,
            patient_context=patient_context,
            allowed_studies=allowed,
            obi_ldl_addon=obi_ldl_combined if topic_key == "ldl_trajectory_trt_apoa2" else "",
            topic_key=topic_key,
        )
        translated = _dehedge_text(translated)
        topic_initials[topic_key] = translated
        v = validate_timeline_attribution(translated, timeline_context)
        topic_validation[topic_key] = v
        if not v.get("clean") and v.get("violations"):
            translated = _dehedge_text(
                _rewrite_timeline_violations(
                    draft_text=translated,
                    timeline_context=timeline_context,
                    violations=v["violations"],
                )
            )
        topic_translations[topic_key] = translated
        cited_pmids.update(_extract_pmids(translated))

    # questions to bring
    discuss_items: set[tuple[str, str]] = set()
    for r in matched:
        if r["action_level"] == "discuss_with_doctor":
            discuss_items.add((_category_display(r["category"]), r["recommendation_text"]))
    for tk, txt in topic_translations.items():
        for item in _extract_discuss_items_from_topic(txt):
            category = "Topic: " + tk
            discuss_items.add((category, item))

    grouped: dict[str, list[str]] = {}
    for cat, item in discuss_items:
        grouped.setdefault(cat, []).append(item)
    for k in grouped:
        grouped[k] = sorted(set(grouped[k]))

    # evidence appendix mapping
    retrieved_by_pmid = {str(s.get("pmid")): s for s in retrieved if s.get("pmid")}
    cited_sorted = sorted([p for p in cited_pmids if p in retrieved_by_pmid], key=lambda x: int(x))

    top5_text = _ai_top_5_actions(
        patient_context_snapshot=snapshot,
        matched_recommendations=matched,
        topic_narratives=topic_translations,
    )
    august_box = _prioritize_august_handoff(grouped, max_items=5)

    bio_db = SessionLocal()
    try:
        mxld = bio_db.query(func.max(LabResult.lab_date)).scalar()
        ref_lab_date: date = mxld if isinstance(mxld, date) else date(2026, 1, 27)
        phen_snapshot = compute_phenotypic_age(bio_db, ref_lab_date)
        fit_snapshot = compute_fitness_age(bio_db, ref_lab_date)
        fram_snapshot = compute_framingham_heart_age(bio_db, ref_lab_date)
    finally:
        bio_db.close()

    # Build docx
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches

    doc = Document()
    configure_document_typography(doc)
    ensure_bullet_styles(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Your Personal Health Report")
    r.bold = True
    r.font.size = Pt(20)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Action-oriented translation of your labs, genetics, and current research - prepared for Robert Grogan")
    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dt.add_run(datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"))
    doc.add_paragraph(
        "This is the companion to the physician report you are bringing to Dr. Lamkin. "
        "It is organized around what you can do now, what you should bring to your next visit, "
        "and what evidence supports each move. Every cited study is real and verifiable by PMID."
    )

    doc.add_heading("Your Top 5 Actions This Month", level=1)
    add_text_with_bullets(doc, top5_text)

    doc.add_heading("Bring to Dr. Lamkin in August", level=2)
    box = doc.add_table(rows=1, cols=1)
    apply_table_layout(box, [6.5])
    box_cell = box.rows[0].cells[0]
    box_cell.text = ""
    for item in august_box:
        box_cell.add_paragraph(f"- {item}", style="List Bullet")

    doc.add_heading("How you're doing right now", level=1)
    _add_biological_age_snapshot_docx(doc, phen_snapshot, fit_snapshot, fram_snapshot)

    for p in section1.split("\n\n"):
        if p.strip():
            doc.add_paragraph(p.strip())

    doc.add_heading("Your genetic playbook", level=1)
    order = ["Diet", "Exercise", "Supplements", "Sleep & Lifestyle", "Monitoring", "Behavior"]
    grouped_recs: dict[str, list[dict[str, Any]]] = {}
    for rec in matched:
        grouped_recs.setdefault(_category_display(rec["category"]), []).append(rec)
    for cat in order:
        entries = grouped_recs.get(cat, [])
        if not entries:
            continue
        doc.add_heading(cat, level=2)
        from docx.shared import Pt, RGBColor

        for e in entries:
            head = doc.add_paragraph()
            head.paragraph_format.space_before = Pt(0)
            head.paragraph_format.space_after = Pt(0)
            rh = head.add_run(e["recommendation_text"])
            rh.bold = True

            body = doc.add_paragraph()
            body.paragraph_format.space_before = Pt(0)
            body.paragraph_format.space_after = Pt(6)
            body.add_run(e["rationale"])
            sep = body.add_run(" — ")
            sep.font.size = Pt(10.5)
            tag_txt = (
                "self-directed"
                if e["action_level"] == "self_directed"
                else "discuss with Dr. Lamkin"
            )
            tag_run = body.add_run(tag_txt)
            tag_run.italic = True
            tag_run.font.size = Pt(9)
            tag_run.font.color.rgb = RGBColor(128, 128, 128)

    for tk, txt in topic_translations.items():
        from app.research.topic_catalog import TOPICS

        title = TOPICS.get(tk, {}).get("title", tk)
        doc.add_heading(title, level=1)
        add_text_with_bullets(doc, txt)

    doc.add_heading("Questions to bring to Dr. Lamkin", level=1)
    for cat in sorted(grouped.keys()):
        doc.add_heading(cat, level=2)
        for item in grouped[cat]:
            doc.add_paragraph(f"- {item}", style="List Bullet")

    doc.add_heading("The evidence behind this", level=1)
    # small appendix table for readability
    tbl = doc.add_table(rows=1, cols=4)
    apply_table_layout(tbl, [0.8, 2.4, 1.0, 2.3])
    headers = ["PMID", "Title", "Journal/Year", "Summary"]
    for i, h in enumerate(headers):
        set_cell_text(tbl.rows[0].cells[i], h, bold=True, size_pt=9)
        shade_cell(tbl.rows[0].cells[i], "D9D9D9")
    for pid in cited_sorted:
        s = retrieved_by_pmid[pid]
        row = tbl.add_row().cells
        vals = [
            pid,
            truncate(s.get("title") or "", 120),
            truncate(f"{s.get('journal','')} ({s.get('year','')})", 70),
            truncate((s.get("summary") or s.get("abstract") or "").replace("\n", " "), 220),
        ]
        for i, v in enumerate(vals):
            set_cell_text(row[i], v, size_pt=9)
            row[i].width = Inches([0.8, 2.4, 1.0, 2.3][i])

    doc.save(str(out_path))

    # update clinical report
    db = SessionLocal()
    try:
        row2 = db.query(ClinicalReport).filter(ClinicalReport.id == report_row.id).first()
        row2.patient_docx_path = str(out_path.resolve())
        db.commit()
    finally:
        db.close()

    # integrity checks summary
    narrative_blob = "\n".join(topic_translations.values())
    hedging_hits = [t for t in HEDGING_TERMS if t in narrative_blob.lower()]
    number_pool = _numbers_from_context(snapshot)
    section1_numbers = _extract_normalized_numbers(section1)
    number_integrity_ok = all(n in number_pool for n in section1_numbers)
    pmid_integrity_ok = all(pid in retrieved_by_pmid for pid in _extract_pmids(narrative_blob))

    # attach debug text fields for requested output extraction
    return {
        "patient_docx_path": str(out_path.resolve()),
        "recommendations_matched": len(matched),
        "topics_translated": len(topic_translations),
        "total_cited_pmids": len(cited_sorted),
        "section1_text": section1,
        "topic_sections": topic_translations,
        "recommendation_samples": random.sample(matched, k=min(3, len(matched))),
        "questions_section": grouped,
        "top5_actions_text": top5_text,
        "august_handoff_items": august_box,
        "hedging_hits": hedging_hits,
        "pmid_integrity_ok": pmid_integrity_ok,
        "number_integrity_ok": number_integrity_ok,
        "timeline_validation": {
            "timeline_context": timeline_context,
            "section1_initial": section1_initial,
            "section1_final": section1,
            "section1_validation": section1_validation,
            "topic_initials": topic_initials,
            "topic_finals": topic_translations,
            "topic_validation": topic_validation,
        },
    }
