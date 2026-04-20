from __future__ import annotations


def truncate(text: str, limit: int) -> str:
    s = (text or "").strip()
    if len(s) <= limit:
        return s
    return s[: max(0, limit - 3)].rstrip() + "..."


def configure_document_typography(doc) -> None:
    from docx.shared import Inches, Pt

    sect = doc.sections[0]
    sect.top_margin = Inches(0.8)
    sect.bottom_margin = Inches(0.8)
    sect.left_margin = Inches(0.8)
    sect.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.line_spacing = 1.15

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.15


def ensure_bullet_styles(doc) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches

    styles = doc.styles
    if "List Bullet" not in styles:
        s = styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Normal"]
        s.paragraph_format.left_indent = Inches(0.25)
        s.paragraph_format.first_line_indent = Inches(-0.15)
    if "List Bullet 2" not in styles:
        s2 = styles.add_style("List Bullet 2", WD_STYLE_TYPE.PARAGRAPH)
        s2.base_style = styles["Normal"]
        s2.paragraph_format.left_indent = Inches(0.5)
        s2.paragraph_format.first_line_indent = Inches(-0.15)


def add_text_with_bullets(doc, text: str) -> None:
    from docx.shared import Inches

    for raw in (text or "").splitlines():
        if not raw.strip():
            doc.add_paragraph("")
            continue
        if raw.startswith("  - "):
            p = doc.add_paragraph(raw[4:].strip(), style="List Bullet 2")
            if p.paragraph_format.left_indent is None:
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.15)
            continue
        if raw.startswith("- "):
            p = doc.add_paragraph(raw[2:].strip(), style="List Bullet")
            if p.paragraph_format.left_indent is None:
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.15)
            continue
        doc.add_paragraph(raw.strip())


def set_cell_padding(cell, twips: int = 80) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(twips))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color: str = "BFBFBF", size: str = "4") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for side in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_cell(cell, fill_hex: str = "D9D9D9") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)


def set_cell_text(cell, text: str, *, bold: bool = False, size_pt: float = 9.0) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt

    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(size_pt)
            r.bold = bold


def apply_table_layout(table, col_widths_in: list[float]) -> None:
    from docx.shared import Inches

    table.autofit = False
    if hasattr(table, "allow_autofit"):
        table.allow_autofit = False

    for i, w in enumerate(col_widths_in):
        table.columns[i].width = Inches(w)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(col_widths_in[i])
            set_cell_padding(cell, 80)
            set_cell_borders(cell, "BFBFBF", "4")
