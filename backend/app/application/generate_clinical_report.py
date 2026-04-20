"""
Generate a doctor-facing clinical DOCX report from PubMed-backed evidence.

No citation may appear unless the study was retrieved from PubMed and graded here.
"""

from __future__ import annotations

import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from sqlalchemy import desc

from app.ai.client import get_openai_client
from app.core.user_context import USER_CONTEXT
from app.db.database import SessionLocal
from app.db.models import ClinicalReport, DailySummary, GeneticVariant
from app.domain.assessment.apple_health_rollup import compute_rollups, format_rollup_block
from app.research.topic_catalog import TOPICS
from app.research.topic_research import research_topic

logger = logging.getLogger(__name__)

DEFAULT_REPORT_TITLE = "Clinical Review for Dr. Brian Lamkin"


def _backend_root() -> Path:
    return Path(__file__).resolve().parents[2]


def _default_output_dir() -> Path:
    d = _backend_root() / "data" / "reports"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _gather_patient_context_bundle(db) -> dict[str, Any]:
    ds = db.query(DailySummary).order_by(desc(DailySummary.created_at)).first()
    genetics = (
        db.query(GeneticVariant)
        .order_by(desc(GeneticVariant.magnitude))
        .limit(20)
        .all()
    )
    rollups = compute_rollups(db)
    phy_text = format_rollup_block(rollups)

    gen_rows = [
        {
            "rsid": g.rsid,
            "genotype": g.genotype,
            "genes": g.genes,
            "magnitude": g.magnitude,
            "summary": g.summary,
        }
        for g in genetics
    ]

    daily_block = {}
    if ds:
        daily_block = {
            "created_at": ds.created_at.isoformat() if ds.created_at else None,
            "summary_text": ds.summary_text,
            "physiology_rollups": ds.physiology_rollups,
            "genetic_context_snippet": ds.genetic_context,
        }

    return {
        "user_context": USER_CONTEXT,
        "latest_daily_summary": daily_block,
        "genetics_top20": gen_rows,
        "physiology_rollups_text": phy_text,
        "physiology_rollups_raw": rollups,
    }


def _patient_context_str(bundle: dict[str, Any]) -> str:
    """Single string for AI prompts."""
    parts = [
        "=== Static user context ===\n" + bundle["user_context"],
        "\n=== Latest cached daily summary ===\n"
        + json.dumps(bundle.get("latest_daily_summary") or {}, indent=2),
        "\n=== Top genetic variants (by magnitude) ===\n"
        + json.dumps(bundle.get("genetics_top20") or [], indent=2),
        "\n=== Physiology rollups (Apple Health) ===\n"
        + (bundle.get("physiology_rollups_text") or ""),
    ]
    return "\n".join(parts)


def _executive_summary_ai(patient_context: str, topic_results: list[dict]) -> str:
    outlines = []
    for tr in topic_results:
        outlines.append(
            f"- {tr['title']}\n  - Question: {tr['research_question']}\n"
            f"  - Narrative excerpt: {(tr.get('narrative') or '')[:1200]}"
        )
    blob = "\n".join(outlines)

    prompt = f"""Write a one-page executive summary preamble for a functional medicine physician (Dr. Brian Lamkin).

Patient context (do not invent clinical facts beyond this):
{patient_context}

Topics researched (narratives are from PubMed-backed synthesis only):
{blob}

Include:
1) Who the patient is in 2-3 sentences (sex, age, location, key therapies only as stated in context).
2) Top 3-5 clinical concerns motivating this packet.
3) What questions the literature sections address.
4) What you are asking the physician to review or discuss.

CRITICAL: Do not invent any findings, lab values, or genetic results not explicitly present in the patient context. Do not cite PMIDs in the executive summary."""

    client = get_openai_client()
    response = client.chat.completions.create(
        model="gpt-5",
        messages=[
            {
                "role": "system",
                "content": "You write concise clinical executive summaries without fabricating data.",
            },
            {"role": "user", "content": prompt},
        ],
    )
    return (response.choices[0].message.content or "").strip()


def _truncate(text: str, limit: int) -> str:
    s = (text or "").strip()
    if len(s) <= limit:
        return s
    return s[: max(0, limit - 3)].rstrip() + "..."


def _configure_document_typography(doc) -> None:
    from docx.shared import Inches, Pt

    sect = doc.sections[0]
    sect.top_margin = Inches(0.8)
    sect.bottom_margin = Inches(0.8)
    sect.left_margin = Inches(0.8)
    sect.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.line_spacing = 1.15

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.15


def _ensure_bullet_styles(doc) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches

    styles = doc.styles
    if "List Bullet" not in styles:
        s = styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Normal"]
        s.paragraph_format.left_indent = Inches(0.25)
        s.paragraph_format.first_line_indent = Inches(-0.15)
    if "List Bullet 2" not in styles:
        s2 = styles.add_style("List Bullet 2", WD_STYLE_TYPE.PARAGRAPH)
        s2.base_style = styles["Normal"]
        s2.paragraph_format.left_indent = Inches(0.5)
        s2.paragraph_format.first_line_indent = Inches(-0.15)


def _add_text_with_bullets(doc, text: str) -> None:
    from docx.shared import Inches

    for raw in (text or "").splitlines():
        if not raw.strip():
            doc.add_paragraph("")
            continue
        if raw.startswith("  - "):
            p = doc.add_paragraph(raw[4:].strip(), style="List Bullet 2")
            if p.paragraph_format.left_indent is None:
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.15)
            continue
        if raw.startswith("- "):
            p = doc.add_paragraph(raw[2:].strip(), style="List Bullet")
            if p.paragraph_format.left_indent is None:
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.15)
            continue
        doc.add_paragraph(raw.strip())


def _set_cell_padding(cell, twips: int = 80) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(twips))
        node.set(qn("w:type"), "dxa")


def _set_cell_borders(cell, color: str = "BFBFBF", size: str = "4") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for side in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _shade_cell(cell, fill_hex: str = "D9D9D9") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)


def _set_cell_text(cell, text: str, *, bold: bool = False, size_pt: float = 9.0) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt

    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(size_pt)
            r.bold = bold


def _apply_table_layout(table, col_widths_in: list[float]) -> None:
    from docx.shared import Inches

    table.autofit = False
    if hasattr(table, "allow_autofit"):
        table.allow_autofit = False

    for i, w in enumerate(col_widths_in):
        table.columns[i].width = Inches(w)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(col_widths_in[i])
            _set_cell_padding(cell, 80)
            _set_cell_borders(cell, "BFBFBF", "4")


def _build_docx(
    *,
    output_path: Path,
    executive_summary: str,
    topic_results: list[dict],
    patient_bundle: dict[str, Any],
) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    _configure_document_typography(doc)
    _ensure_bullet_styles(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Clinical Review for Dr. Brian Lamkin")
    r.bold = True
    r.font.size = Pt(18)

    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Patient: male, 54 - prepared from patient-entered data and PubMed literature")

    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_ts = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Generated: {gen_ts}\n")
    p3.add_run("Prepared using AI-assisted synthesis of patient data and PubMed literature.")

    doc.add_page_break()

    doc.add_heading("Disclaimer", level=1)
    tbl_d = doc.add_table(rows=1, cols=1)
    _apply_table_layout(tbl_d, [6.5])
    disc = (
        "This report is generated by AI from patient-provided data and public medical literature. "
        "It is intended to support, not replace, clinical judgment. All cited studies are retrieved from PubMed. "
        "AI-assigned relevance and quality grades are based on abstract review only. "
        "Please verify any citation of clinical significance."
    )
    _set_cell_text(tbl_d.rows[0].cells[0], disc, size_pt=10)

    doc.add_heading("Executive Summary", level=1)
    _add_text_with_bullets(doc, executive_summary)

    for tr in topic_results:
        doc.add_heading(tr["title"], level=1)
        doc.add_paragraph(f"Research question: {tr['research_question']}")
        doc.add_heading("Synthesis", level=2)
        _add_text_with_bullets(doc, tr.get("narrative") or "")

        doc.add_heading("Evidence Table", level=2)
        studies = tr.get("studies") or []
        if not studies:
            doc.add_paragraph("No studies met inclusion criteria after grading.")
            continue

        table = doc.add_table(rows=1, cols=9)
        col_widths = [0.6, 1.4, 0.9, 0.7, 0.3, 0.3, 0.4, 1.2, 0.7]
        _apply_table_layout(table, col_widths)

        headers = [
            "PMID",
            "Title",
            "Journal / Year",
            "Types",
            "Rel",
            "Imp",
            "Conf",
            "Summary",
            "Applicability",
        ]
        for i, h in enumerate(headers):
            _set_cell_text(table.rows[0].cells[i], h, bold=True, size_pt=9)
            _shade_cell(table.rows[0].cells[i], "D9D9D9")

        for s in studies:
            row = table.add_row().cells
            jy = f"{s.get('journal', '')} ({s.get('year', '')})"
            stypes = ", ".join(s.get("publication_types") or [])
            vals = [
                str(s.get("pmid", "")),
                _truncate(s.get("title") or "", 90),
                _truncate(jy, 90),
                _truncate(stypes, 80),
                str(s.get("relevance_grade", "")),
                str(s.get("importance_grade", "")),
                str(s.get("confidence", "")),
                _truncate(s.get("summary") or "", 400),
                _truncate(s.get("applicability_note") or "", 250),
            ]
            for i, v in enumerate(vals):
                _set_cell_text(row[i], v, size_pt=9)
                row[i].width = Inches(col_widths[i])

    doc.add_page_break()
    doc.add_heading("Appendix A - Current regimen (from patient context)", level=1)
    _add_text_with_bullets(doc, _extract_regimen_block(USER_CONTEXT))

    doc.add_heading("Appendix B - Genetic variants (top 20 by magnitude)", level=1)
    for g in patient_bundle.get("genetics_top20") or []:
        doc.add_paragraph(
            f"{g.get('rsid')} {g.get('genotype')} - genes: {g.get('genes')} - mag={g.get('magnitude')} - {g.get('summary')}"
        )

    doc.add_heading("Appendix C - Latest lab snapshot (cached summary)", level=1)
    lt = patient_bundle.get("latest_daily_summary") or {}
    _add_text_with_bullets(
        doc, lt.get("summary_text") or "(none - run daily summary generation first)"
    )

    doc.add_heading("Appendix D - Physiology rollups", level=1)
    _add_text_with_bullets(doc, patient_bundle.get("physiology_rollups_text") or "")
    doc.save(str(output_path))


def _extract_regimen_block(uc: str) -> str:
    lines = uc.splitlines()
    out: list[str] = []
    capture = False
    for ln in lines:
        if ln.strip().startswith("Current prescription therapies:"):
            capture = True
        if capture:
            out.append(ln)
        if ln.strip().startswith("Known genetic predispositions"):
            break
    return "\n".join(out).strip() or uc


def generate_clinical_report(
    topic_keys: list[str] | None = None,
    output_dir: str | None = None,
) -> dict[str, Any]:
    keys = topic_keys if topic_keys is not None else list(TOPICS.keys())
    for k in keys:
        if k not in TOPICS:
            raise ValueError(f"Unknown topic key: {k}")

    out_dir = Path(output_dir) if output_dir else _default_output_dir()
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    docx_path = out_dir / f"clinical_report_{ts}.docx"

    db = SessionLocal()
    try:
        bundle = _gather_patient_context_bundle(db)
    finally:
        db.close()

    patient_context = _patient_context_str(bundle)
    topic_results: list[dict] = []
    all_kept: list[dict] = []
    retrieved_total = 0

    for key in keys:
        cfg = TOPICS[key]
        try:
            res = research_topic(cfg, patient_context)
            res["topic_key"] = key
            topic_results.append(res)
            retrieved_total += int(res.get("abstracts_fetched") or 0)
            for s in res.get("studies") or []:
                row = dict(s)
                row["topic_key"] = key
                all_kept.append(row)
        except Exception as e:
            logger.exception("Topic %s failed: %s", key, e)
            topic_results.append(
                {
                    "topic_key": key,
                    "title": cfg["title"],
                    "research_question": cfg["research_question"],
                    "studies": [],
                    "abstracts_fetched": 0,
                    "narrative": f"Topic generation failed: {e}",
                }
            )

    executive = _executive_summary_ai(patient_context, topic_results)
    narrative_sections = {
        str(tr.get("topic_key") or ""): tr.get("narrative", "") for tr in topic_results
    }

    _build_docx(
        output_path=docx_path,
        executive_summary=executive,
        topic_results=topic_results,
        patient_bundle=bundle,
    )

    snapshot = json.dumps(bundle, default=str)
    studies_json = json.dumps(all_kept, default=str)
    narrative_json = json.dumps(narrative_sections, ensure_ascii=False)
    topics_json = json.dumps([TOPICS[k]["title"] for k in keys], ensure_ascii=False)

    db = SessionLocal()
    try:
        row = ClinicalReport(
            report_title=DEFAULT_REPORT_TITLE,
            topic_areas=topics_json,
            patient_context_snapshot=snapshot,
            retrieved_studies=studies_json,
            narrative_sections=narrative_json,
            executive_summary=executive,
            docx_path=str(docx_path.resolve()),
        )
        db.add(row)
        db.commit()
        db.refresh(row)
        rid = row.id
    finally:
        db.close()

    return {
        "report_id": rid,
        "docx_path": str(docx_path.resolve()),
        "topic_count": len(keys),
        "studies_retrieved_count": retrieved_total,
        "studies_kept_count": len(all_kept),
    }
